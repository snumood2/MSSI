import json
import re
import subprocess
from difflib import SequenceMatcher
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
DOCX = Path("/Users/wm/Library/CloudStorage/OneDrive-개인/9. Temporal Work/2026 설문웹앱개밸/기분장애_설문지.docx")


def norm(text):
    text = re.sub(r"\[cite:[^\]]+\]", "", str(text or ""))
    text = re.sub(r"^[\s>]*(?:[A-Z]\d+[a-z]?|[A-Z]\.|[a-z]\)|\d+(?:\.\d+)?\.?)\s*", "", text)
    text = re.sub(r"\s+", "", text)
    text = text.replace("'", "").replace('"', "").replace("‘", "").replace("’", "")
    text = text.replace("“", "").replace("”", "").replace("ㆍ", "/")
    return text


def ratio(a, b):
    a = norm(a)
    b = norm(b)
    if not a and not b:
        return 1.0
    return SequenceMatcher(None, a, b).ratio()


def option_ratio(doc_options, web_options):
    def clean_option(item):
        return norm(re.sub(r"^\s*\d+\s*[:.]\s*", "", str(item or "")))
    return ratio(" ".join(clean_option(x) for x in doc_options), " ".join(clean_option(x) for x in web_options))


def web_sections():
    code = """
import { SURVEY_SECTIONS } from './questions.js';
console.log(JSON.stringify(SURVEY_SECTIONS));
"""
    out = subprocess.check_output(
        ["node", "--input-type=module", "-e", code],
        cwd=ROOT,
        text=True,
    )
    return json.loads(out)


def clean_cell(cell):
    return re.sub(r"\s+", " ", cell.text).strip()


def doc_table_questions(table, start_row=1, q_col=1):
    rows = []
    for row in table.rows[start_row:]:
        cells = [clean_cell(c) for c in row.cells]
        if len(cells) <= q_col:
            continue
        q = cells[q_col]
        if not q or q in {
            "문 항",
            "평소에 나는……",
            "평소에…",
            "어렸을 때 나는……",
            "어렸을 때 학교에서 나는……",
            "지난 주에 나는…",
        }:
            continue
        rows.append(q)
    return rows


def mssi_questions(table):
    rows = []
    for row in table.rows[1:]:
        first = clean_cell(row.cells[0])
        if re.match(r"^\d+\.\s", first):
            rows.append(first)
    return rows


def doc_table_options(table):
    if not table.rows:
        return []
    cells = [clean_cell(c) for c in table.rows[0].cells]
    return [c for c in cells[2:] if c]


def web_questions(section):
    return [q.get("text", "") for q in section.get("questions", []) if q.get("type") != "info"]


def web_options(section):
    opts = section.get("options") or []
    if opts and isinstance(opts[0], dict) and "opts" in opts[0]:
        return [f"{group.get('label')}: " + " / ".join(o.get("l", "") for o in group.get("opts", [])) for group in opts]
    if opts:
        return [o.get("l", "") for o in opts]
    return []


def option_texts_from_questions(section):
    values = []
    for q in section.get("questions", []):
        opts = q.get("options") or []
        if opts:
            values.extend(o.get("l", "") for o in opts)
    return values


TABLE_MAP = [
    ("zung_sds", 1, 1, 1),
    ("bai", 2, 1, 1),
    ("temps_a", 3, 1, 1),
    ("mssi", 4, 1, 0),
    ("miq_t", 11, 1, 1),
    ("k_mdq", 12, 1, 1),
    ("bapq", 13, 1, 1),
    ("ctq_sf", 14, 1, 1),
    ("ipsm", 15, 1, 1),
    ("cd_risc", 16, 2, 1),
    ("ersq", 17, 2, 1),
    ("bis_bas", 18, 2, 1),
    ("audit_k", 20, 1, 1),
    ("asrs", 40, 1, 1),
    ("pai_bor", 41, 1, 1),
    ("wurs", 42, 2, 1),
]


def main():
    doc = Document(DOCX)
    sections = {s["id"]: s for s in web_sections()}
    results = []

    for section_id, table_no, start_row, q_col in TABLE_MAP:
        section = sections[section_id]
        table = doc.tables[table_no - 1]
        if section_id == "mssi":
            doc_q = mssi_questions(table)
        elif section_id == "bis_bas":
            doc_q = doc_table_questions(doc.tables[17], 2, 1) + doc_table_questions(doc.tables[18], 2, 1)
        else:
            doc_q = doc_table_questions(table, start_row=start_row, q_col=q_col)
        if section_id == "k_mdq":
            doc_q = doc_q[:15]
        web_q = web_questions(section)
        mismatches = []
        for i in range(max(len(doc_q), len(web_q))):
            dq = doc_q[i] if i < len(doc_q) else ""
            wq = web_q[i] if i < len(web_q) else ""
            r = ratio(dq, wq)
            if r < 0.985:
                mismatches.append({
                    "index": i + 1,
                    "similarity": round(r, 3),
                    "docx": dq,
                    "web": wq,
                })
        opt_doc = doc_table_options(table)
        opt_web = web_options(section)
        option_note = ""
        if section_id == "mssi":
            opt_doc = [
                "1 = 1일 미만 (몇 시간 지속) / 2 = 1일 이상 1주 미만 / 3 = 1주 이상 2주 미만 / 4 = 매일 증상이 있었다",
                "1 = 약간 (증상은 있었지만 힘들지 않았다) / 2 = 상당히 (매우 불쾌하지만 참을 수 있었다) / 3 = 심하게 (정말 견디기 힘들었다)",
            ]
            opt_web = web_options(section)
            option_note = "MSSI frequency/severity options compared from DOCX instruction text."
        elif section_id == "k_mdq":
            opt_doc = ["1. 예", "0. 아니오", "0. 문제 없었다", "1. 경미한 문제", "2. 중등도의 문제", "3. 심각한 문제"]
            opt_web = ["1. 예", "0. 아니오"] + option_texts_from_questions(section)[-4:]
            option_note = "MDQ yes/no and impairment options compared from split DOCX rows."
        elif section_id == "audit_k":
            opt_doc = []
            for row in table.rows[1:]:
                cells = [clean_cell(c) for c in row.cells]
                opt_doc.extend(c for c in cells[2:] if c)
            opt_web = option_texts_from_questions(section)
            option_note = "AUDIT-K per-question options compared from each row."
        opt_ratio = option_ratio(opt_doc, opt_web)
        results.append({
            "section_id": section_id,
            "title": section["title"],
            "docx_table": table_no,
            "docx_questions": len(doc_q),
            "web_questions": len(web_q),
            "question_mismatch_count": len(mismatches),
            "question_mismatches": mismatches[:30],
            "docx_options": opt_doc,
            "web_options": opt_web,
            "option_similarity": round(opt_ratio, 3),
            "option_note": option_note,
        })

    # Composite sections not represented by one simple matrix table.
    csm = sections["csm"]
    csm_doc_questions = []
    for line in Path("/tmp/mssi_docx_text.txt").read_text(encoding="utf-8").splitlines():
        if re.match(r"^(하루의 할 일을|저녁시간에|평소 아침에|아침에 깨어난|친구와 함께|귀하는 저녁에|중요한 시험|사람에게는 아침형|하루 8시간|항상 아침 6시|밤잠을 자고|당신은 아침이나)", line):
            csm_doc_questions.append(line)
    csm_web = web_questions(csm)
    csm_mismatches = []
    for i in range(max(len(csm_doc_questions), len(csm_web))):
        dq = csm_doc_questions[i] if i < len(csm_doc_questions) else ""
        wq = csm_web[i] if i < len(csm_web) else ""
        r = ratio(dq, wq)
        if r < 0.985:
            csm_mismatches.append({"index": i + 1, "similarity": round(r, 3), "docx": dq, "web": wq})
    results.append({
        "section_id": "csm",
        "title": csm["title"],
        "docx_table": "21-33 plus prompts",
        "docx_questions": len(csm_doc_questions),
        "web_questions": len(csm_web),
        "question_mismatch_count": len(csm_mismatches),
        "question_mismatches": csm_mismatches[:30],
        "option_similarity": None,
    })

    pms = sections["pms"]
    pms_doc_q = doc_table_questions(doc.tables[42], 1, 1) + doc_table_questions(doc.tables[43], 1, 1)
    pms_web = web_questions(pms)
    pms_mismatches = []
    for i in range(max(len(pms_doc_q), len(pms_web))):
        dq = pms_doc_q[i] if i < len(pms_doc_q) else ""
        wq = pms_web[i] if i < len(pms_web) else ""
        r = ratio(dq, wq)
        if r < 0.985:
            pms_mismatches.append({"index": i + 1, "similarity": round(r, 3), "docx": dq, "web": wq})
    results.append({
        "section_id": "pms",
        "title": pms["title"],
        "docx_table": "43-44",
        "docx_questions": len(pms_doc_q),
        "web_questions": len(pms_web),
        "question_mismatch_count": len(pms_mismatches),
        "question_mismatches": pms_mismatches[:30],
        "option_similarity": option_ratio(doc_table_options(doc.tables[42]), web_options(pms)),
    })

    out = {
        "docx": str(DOCX),
        "summary": {
            "sections_checked": len(results),
            "sections_with_question_mismatches": sum(1 for r in results if r["question_mismatch_count"]),
            "total_question_mismatches": sum(r["question_mismatch_count"] for r in results),
            "sections_with_option_drift": [
                r["section_id"] for r in results
                if r.get("option_similarity") is not None and r["option_similarity"] < 0.9
            ],
        },
        "results": results,
    }
    out_path = ROOT / "docx_question_audit_report.json"
    out_path.write_text(json.dumps(out, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(out["summary"], ensure_ascii=False, indent=2))
    print(out_path)


if __name__ == "__main__":
    main()
